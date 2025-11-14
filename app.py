import os
import openai
import json
import time
import streamlit as st
from docx import Document
import io
import fitz
from supabase import create_client, Client

# --- 0. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(
    page_title="Olea Asistente Legal", 
    layout="centered",
    page_icon="favicon.png"
)

st.markdown(
    """
    <meta property="og:title" content="Olea Abogados - Asistente Legal IA">
    <meta property="og:description" content="Herramienta de IA para la generación y análisis de documentos legales.">
    <meta property="og:image" content="URL_PUBLICA_DE_TU_LOGO_AQUI">
    """,
    unsafe_allow_html=True
)

# --- ¡NUEVO BLOQUE DE CREDENCIALES V2.3 (Doble Cliente)! ---
OPENAI_API_KEY_VAL = None
SUPABASE_URL_VAL = None
SUPABASE_KEY_VAL = None
SUPABASE_SERVICE_KEY_VAL = None # <-- ¡NUEVO!

# Leer todas las llaves (Streamlit Cloud o Servidor de Nuria)
try:
    OPENAI_API_KEY_VAL = st.secrets["OPENAI_API_KEY"]
    SUPABASE_URL_VAL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY_VAL = st.secrets["SUPABASE_KEY"]
    SUPABASE_SERVICE_KEY_VAL = st.secrets["SUPABASE_SERVICE_KEY"] # <-- ¡NUEVO!

except (KeyError, FileNotFoundError):
    OPENAI_API_KEY_VAL = os.environ.get("OPENAI_API_KEY")
    SUPABASE_URL_VAL = os.environ.get("SUPABASE_URL")
    SUPABASE_KEY_VAL = os.environ.get("SUPABASE_KEY")
    SUPABASE_SERVICE_KEY_VAL = os.environ.get("SUPABASE_SERVICE_KEY") # <-- ¡NUEVO!

# 1. Conectar a OpenAI (Sin cambios)
try:
    client = openai.OpenAI(api_key=OPENAI_API_KEY_VAL)
except Exception as e:
    st.error(f"Error fatal: No se pudo conectar a OpenAI. Revisa la API Key.")
    st.stop()

# 2. Conectar a Supabase (¡DOBLE CONEXIÓN!)
try:
    # Cliente ANÓNIMO (para el Login de todos)
    supabase_anon: Client = create_client(SUPABASE_URL_VAL, SUPABASE_KEY_VAL)
    
    # Cliente ADMIN (¡LA LLAVE MAESTRA! Solo para el CRUD)
    supabase_admin: Client = create_client(SUPABASE_URL_VAL, SUPABASE_SERVICE_KEY_VAL)
except Exception as e:
    st.error(f"Error fatal: No se pudo conectar a Supabase. Revisa las llaves de Supabase.")
    st.stop()

# --- FIN DEL BLOQUE NUEVO ---

ADMIN_EMAIL = ["alejandro@neurya.com", "elisa@oleaabogados.com"]

# --- 2. CONFIGURACIÓN GLOBAL Y COMPONENTES 1, 2, 3, 4 ---
# (Todo nuestro código de "motor" va aquí, definido como funciones)

# Conectar a OpenAI
try:
    client = openai.OpenAI()
except Exception as e:
    # Esto pasará si la API key no está configurada
    st.error(f"Error al conectar con OpenAI. ¿Configuraste la variable OPENAI_API_KEY?")
    st.stop() # Detiene la app si no hay API key

# Constantes del Template
NOMBRES_DE_ESTILOS = ['Titulo_1', 'Parrafo_Justificado', 'Lista_Numerada', 'Estilo_Firma', 'Lista_Manual']
TEMPLATE_FILE_GENERAL = 'template_maestro.docx'
TEMPLATE_FILE_PAGARE = 'template_pagare.docx'

# --- COMPONENTE 2: El "Extractor Universal" (V1.9) ---
@st.cache_data
def extraer_texto_del_documento(archivo_subido):
    """
    Extractor inteligente V1.9: Lee .docx y .pdf
    Toma un archivo subido por Streamlit (UploadedFile)
    y devuelve una sola cadena de texto.
    """
    try:
        # 1. Leer el archivo en memoria una sola vez
        file_bytes = archivo_subido.read()
        
        # 2. Decidir qué herramienta usar
        if archivo_subido.type == "application/pdf":
            # --- LÓGICA PARA PDF (NUEVO) ---
            st.write("Detectado: PDF. Usando PyMuPDF (fitz)...")
            texto_completo = []
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    texto_completo.append(page.get_text())
            return '\n'.join(texto_completo)

        elif archivo_subido.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # --- LÓGICA PARA DOCX (LA ANTIGUA) ---
            st.write("Detectado: DOCX. Usando python-docx...")
            doc = Document(io.BytesIO(file_bytes))
            texto_completo = []
            for parrafo in doc.paragraphs:
                if parrafo.text.strip():
                    texto_completo.append(parrafo.text)
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        if celda.text.strip():
                            texto_completo.append(celda.text)
            return '\n'.join(texto_completo)
        
        else:
            st.error(f"Tipo de archivo no soportado: {archivo_subido.type}")
            return None
            
    except Exception as e:
        st.error(f"Error al leer el archivo ({archivo_subido.name}): {e}")
        return None

@st.cache_data
def generar_documento_ia_general(instruccion_usuario, texto_de_ejemplo, modelo_ia):
    st.info(f"Conectando con la IA ({modelo_ia})... Esto puede tardar varios minutos.")
    lista_estilos_str = ", ".join([f"'{s}'" for s in NOMBRES_DE_ESTILOS])
    prompt_final = f"""
    PRIORIDAD ABSOLUTA: Tu única y exclusiva respuesta debe ser un objeto JSON válido.
    Tu respuesta debe comenzar con un corchete de apertura [ y terminar con un corchete de cierre ].
    No escribas NADA antes del [ ni NADA después del ].
    Eres un abogado senior 'todoterreno' en México.
    FORMATO DE SALIDA OBLIGATORIO:
    Una lista de objetos JSON: [{{"style": "...", "text": "..."}}]
    REGLAS DE ESTILO OBLIGATORIAS:
    - Los únicos valores permitidos para "style" son: {lista_estilos_str}.
    - Usa 'Titulo_1' para todos los títulos.
    - Usa 'Parrafo_Justificado' para el texto principal.
    - REGLAS DE LISTAS:
        - Para cláusulas largas que deben continuar (ej. 'CLÁUSULA PRIMERA', 'CLÁUSULA SEGUNDA'), usa el estilo 'Lista_Numerada'.
        - Para cualquier OTRA lista (ej. un Orden del Día, listas de requisitos, etc.) que deba empezar en '1.', usa el estilo 'Lista_Manual'.
        - IMPORTANTE: Cuando uses 'Lista_Manual', DEBES escribir tú mismo el número y un tabulador. Ejemplo: "1.\t[Texto del punto uno]", "2.\t[Texto del punto dos]", etc.
    - Usa 'Estilo_Firma' para las firmas.
    EJEMPLOS DE TONO (PARA IMITAR):
    ---
    {texto_de_ejemplo}
    ---
    [TAREA DEL USUARIO]:
    {instruccion_usuario}
    Recuerda: Tu respuesta debe ser solo el JSON, empezando con [ y terminando con ].
    """
    try:
        start_time = time.time()
        completion = client.responses.create(
            model=modelo_ia, # Usamos el modelo seleccionado
            input=prompt_final,
            max_output_tokens=20000
        )
        end_time = time.time()
        st.info(f"Respuesta recibida de la IA en {end_time - start_time:.2f} segundos.")
        
        respuesta_cruda = completion.output_text
        start_index = respuesta_cruda.find('[')
        end_index = respuesta_cruda.rfind(']')
        
        if start_index == -1 or end_index == -1:
            st.error("Error: La IA no devolvió un JSON válido. Respuesta cruda:")
            st.code(respuesta_cruda)
            return None
            
        json_limpio = respuesta_cruda[start_index : end_index + 1]
        return json_limpio
        
    except Exception as e:
        st.error(f"ERROR INESPERADO DE API: {e}")
        return None
    
# --- CEREBRO 2.0 (PAGARÉ) - V2.0 (CON CÁLCULO DE IVA) ---
@st.cache_data
def generar_pagare_ia(instruccion_usuario, texto_de_ejemplo, modelo_ia):
    st.info(f"Conectando con la IA ({modelo_ia}) para el Pagaré [ChatEndpoint]...")
    
    lista_estilos_str = ", ".join([f"'{s}'" for s in NOMBRES_DE_ESTILOS])
    
    # --- ¡PROMPT V2.0 ACTUALIZADO CON IVA! ---
    prompt_final = f"""
    PRIORIDAD ABSOLUTA: Tu única y exclusiva respuesta debe ser un objeto JSON válido.
    Tu respuesta debe comenzar con un corchete de apertura {{ y terminar con un corchete de cierre }}.
    No escribas NADA antes del {{ ni NADA después del }}.

    Eres un abogado y contador experto en México.
    
    FORMATO DE SALIDA OBLIGATORIO:
    Un objeto JSON con dos claves principales: "prosa" y "tabla_amortizacion".

    1. CLAVE "prosa":
       Una lista de objetos JSON para el texto legal del pagaré.
       "prosa": [{{"style": "...", "text": "..."}}]
       (Usa los estilos: {lista_estilos_str})

    2. CLAVE "tabla_amortizacion":
       Una lista de objetos JSON con los datos de cada pago de la tabla de amortización.
       
       ¡TAREA DE CÁLCULO CLAVE!:
       Debes calcular los intereses Y ADEMÁS el 16% de IVA sobre dichos intereses.
       El "Pago Total" será la suma de (Capital + Interés + IVA del Interés).

       FORMATO JSON DE TABLA (6 Columnas):
       "tabla_amortizacion": [
         {{"Pago No.": 1, "Interés": 100.00, "IVA del Interés": 16.00, "Capital": 900.00, "Pago Total": 1016.00, "Saldo Insoluto": 9900.00}}
       ]
       
       REGLAS DE FORMATO PARA "tabla_amortizacion":
       - Las claves y el ORDEN deben ser exactamente: "Pago No.", "Interés", "IVA del Interés", "Capital", "Pago Total", "Saldo Insoluto".
       - Los valores deben ser números (float para los montos).

    EJEMPLOS DE TONO (PARA IMITAR):
    ---
    {texto_de_ejemplo}
    ---
    
    [TAREA DEL USUARIO]:
    {instruccion_usuario}
    
    Recuerda: Tu respuesta debe ser solo el JSON, empezando con {{ y terminando con }}.
    """
    
    try:
        start_time = time.time()
        
        completion = client.chat.completions.create(
            model=modelo_ia, 
            messages=[
                {"role": "system", "content": prompt_final},
                {"role": "user", "content": instruccion_usuario}
            ],
            response_format={"type": "json_object"}, 
            max_completion_tokens=20000
        )
        end_time = time.time()
        st.info(f"Respuesta recibida de la IA en {end_time - start_time:.2f} segundos.")
        
        json_respuesta = completion.choices[0].message.content
        return json_respuesta
        
    except Exception as e:
        st.error(f"ERROR INESPERADO DE API (Pagaré V2.0): {e}")
        return None

# --- ENSAMBLADOR 1.0 (GENERAL) - CORREGIDO V1.2 ---
def ensamblar_docx_general(json_data, archivo_template): # <--- CAMBIO 1: Acepta el 2do argumento
    """
    Toma el JSON de la IA y construye un .docx
    usando el template general.
    """
    try:
        # --- CAMBIO 2: Usamos la variable 'archivo_template' ---
        if not os.path.exists(archivo_template): 
            st.error(f"¡Error crítico! No se encuentra el archivo '{archivo_template}'.")
            return None
            
        doc = Document(archivo_template) # <--- CAMBIO 3: Usamos la variable
        # --- FIN DE CAMBIOS ---
        
        # (El resto del código ya estaba bien)
        datos = json.loads(json_data)
        
        for item in datos:
            texto = item.get('text', '')
            estilo = item.get('style', 'Parrafo_Justificado') 
            
            if estilo not in NOMBRES_DE_ESTILOS:
                # Usamos st.warning para que se vea en la app
                st.warning(f"Advertencia: La IA usó un estilo desconocido ('{estilo}'). Usando 'Parrafo_Justificado'.")
                estilo = 'Parrafo_Justificado'
            
            doc.add_paragraph(texto, style=estilo)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except json.JSONDecodeError:
        st.error(f"ERROR CRÍTICO: La IA no devolvió un JSON válido. No se puede ensamblar.")
        st.code(json_data) # Muestra el JSON roto
        return None
    except Exception as e:
        st.error(f"ERROR durante el ensamblaje del .docx: {e}")
        return None
    
# --- ENSAMBLADOR 2.0 (PAGARÉ) - V2.0 (Dibuja 6 columnas con IVA) ---
def ensamblar_pagare_en_memoria(json_data, archivo_template):
    """
    Toma el JSON de pagaré (prosa y tabla V2.0 con IVA) 
    y construye un .docx EN MEMORIA.
    Esta versión CREA la tabla desde cero al final del documento.
    """
    try:
        if not os.path.exists(archivo_template):
            st.error(f"¡Error crítico! No se encuentra el archivo '{archivo_template}'.")
            return None
            
        doc = Document(archivo_template) 
        datos = json.loads(json_data)
        
        # --- 1. Ensamblar la Prosa (Sin cambios) ---
        st.write("Ensamblando prosa del pagaré...")
        prosa_items = datos.get("prosa", [])
        for item in prosa_items:
            texto = item.get('text', '')
            estilo = item.get('style', 'Parrafo_Justificado')
            if estilo not in NOMBRES_DE_ESTILOS:
                estilo = 'Parrafo_Justificado'
            doc.add_paragraph(texto, style=estilo)
            
        # --- 2. Ensamblar la Tabla de Amortización (¡CON 6 COLUMNAS!) ---
        tabla_datos = datos.get("tabla_amortizacion", [])
        if tabla_datos:
            st.write("Creando y ensamblando tabla de amortización (con IVA)...")
            
            # --- ¡CAMBIO 1: Creamos 6 columnas! ---
            table = doc.add_table(rows=1, cols=6) 
            table.style = 'Table Grid'
            
            # --- ¡CAMBIO 2: Rellenamos 6 encabezados! ---
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Pago No.'
            hdr_cells[1].text = 'Interés'
            hdr_cells[2].text = 'IVA del Interés' # <-- NUEVO
            hdr_cells[3].text = 'Capital'
            hdr_cells[4].text = 'Pago Total'      # <-- NUEVO
            hdr_cells[5].text = 'Saldo Insoluto'
            
            # (Opcional: Poner encabezados en negrita)
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True

            # --- ¡CAMBIO 3: Rellenamos 6 celdas! ---
            for fila_data in tabla_datos:
                row_cells = table.add_row().cells # Añade una nueva fila
                
                # Usamos la V1.8 (orden) que es más segura
                valores = list(fila_data.values())
                
                try:
                    row_cells[0].text = str(valores[0]) # Pago No.
                    row_cells[1].text = f"{float(valores[1]):,.2f}" # Interés
                    row_cells[2].text = f"{float(valores[2]):,.2f}" # IVA (Nuevo)
                    row_cells[3].text = f"{float(valores[3]):,.2f}" # Capital
                    row_cells[4].text = f"{float(valores[4]):,.2f}" # Pago Total (Nuevo)
                    row_cells[5].text = f"{float(valores[5]):,.2f}" # Saldo
                except Exception as e:
                    st.warning(f"Error al procesar fila de tabla: {e}. Datos: {valores}")
                    row_cells[0].text = "ERROR"
            
            doc.add_paragraph("") 
                
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Error durante el ensamblaje del pagaré .docx: {e}")
        return None

# --- 3. FUNCIÓN PARA MOSTRAR LA APP PRINCIPAL ---
# (Todo nuestro Componente 5 va aquí adentro)
def mostrar_app_principal():
    
    # Botón de Cerrar Sesión (¡Nuevo!)
    st.sidebar.button("Cerrar Sesión", on_click=logout)

    # Logo
    st.sidebar.image("logo.png")

    st.sidebar.markdown("---")
    st.sidebar.markdown("### Configuración de IA")
    modelo_seleccionado = st.sidebar.selectbox(
        "Elige el motor de IA:",
        ("gpt-5", "gpt-5-pro"),
        index=0, 
        help="GPT-5 es más rápido y barato. GPT-5-Pro es más lento, caro, pero más inteligente."
    )
    st.sidebar.caption(f"Usando: {modelo_seleccionado}")
    
    # Guardamos el modelo en la memoria para que el chat lo vea
    st.session_state['modelo_seleccionado'] = modelo_seleccionado

    # (Dentro de def mostrar_app_principal():)
    
    # --- ¡CAMBIO! Lógica de Pestañas Dinámica (V2.3) ---
    st.title("🤖 Asistente Legal")
    
    lista_pestañas = ["Generador de Documentos", "Chatbot Legal (Analizador)"]
    
    # Revisar si el usuario logueado es el Admin
    if st.session_state.get('user_email') in ADMIN_EMAIL:
        lista_pestañas.append("👑 Administración") # <-- Añade la pestaña Admin
    
    tabs = st.tabs(lista_pestañas)
    
    # Asignar contenido a las pestañas
    with tabs[0]:
        mostrar_pagina_generador(modelo_seleccionado) 

    with tabs[1]:
        mostrar_pagina_chatbot() 
        
    # Si la pestaña Admin existe, asignarle su función
    if len(tabs) == 3:
        with tabs[2]:
            mostrar_pagina_admin()
            
    # --- FIN DEL CAMBIO ---

# --- 3. FUNCIONES DE WORKFLOW ---

def ejecutar_flujo_general(instruccion, ejemplo_subido, modelo_seleccionado):
    """
    Este es el flujo de trabajo que YA CONSTRUIMOS.
    """
    st.success("Iniciando flujo: Documento General")
    
    texto_de_ejemplo = "N/A"
    if ejemplo_subido:
        with st.spinner("Leyendo documento de ejemplo..."):
            texto_de_ejemplo = extraer_texto_del_documento(ejemplo_subido)
    
    if texto_de_ejemplo is not None:
        with st.spinner(f"🧠✨ La IA ({modelo_seleccionado}) está redactando... Esto puede tardar varios minutos."):
            # ¡IMPORTANTE! Tenemos que definir qué prompt usar.
            # Por ahora, solo tenemos uno.
            json_respuesta = generar_documento_ia_general(instruccion, texto_de_ejemplo, modelo_seleccionado)
        
        if json_respuesta:
            with st.spinner("🛠️ Ensamblando el archivo .docx..."):
                # ¡IMPORTANTE! Le decimos qué template usar
                buffer_docx = ensamblar_docx_general(json_respuesta, TEMPLATE_FILE_GENERAL)
            
            if buffer_docx:
                st.success("¡Tu documento está listo para descargar!")
                st.download_button(
                    label="📥 Descargar Documento Generado",
                    data=buffer_docx,
                    file_name="DOCUMENTO_GENERADO.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

def mostrar_pagina_generador(modelo_seleccionado):
    st.markdown("Genera borradores de documentos legales usando IA.")

    st.markdown("### 1. Selecciona el tipo de documento")
    tipo_doc = st.selectbox(
        "Elige la plantilla que deseas usar:",
        ("Documento General", "Pagaré (con Tabla)")
    )

    st.markdown("---") # Una línea divisoria

    instruccion = st.text_area(
        "2. Escribe las instrucciones del documento que necesitas:",
        height=200,
        placeholder="Ej: Genera un Contrato de Arrendamiento simple para una casa. Arrendador: Juan Pérez. Arrendatario: María López..."
    )

    ejemplo_subido = st.file_uploader(
        "3. (Opcional) Sube un .docx o pfd de ejemplo para imitar el tono:",
        type=["docx", "pdf"]
    )

    if st.button("🚀 Generar Documento"):
        if not instruccion:
            st.warning("Por favor, escribe las instrucciones del documento.")
        else:
            # --- IF para decidir qué flujo usar ---
            if tipo_doc == "Documento General":
                ejecutar_flujo_general(instruccion, ejemplo_subido, modelo_seleccionado)
            
            elif tipo_doc == "Pagaré (con Tabla)":
                st.success("Iniciando flujo: Pagaré (con Tabla)")
                texto_de_ejemplo = "N/A"
                if ejemplo_subido:
                    with st.spinner("Leyendo documento de ejemplo..."):
                        texto_de_ejemplo = extraer_texto_del_documento(ejemplo_subido)
                
                if texto_de_ejemplo is not None:
                    with st.spinner(f"🧠✨ La IA ({modelo_seleccionado}) está redactando y calculando..."):
                        # Usamos la nueva función del cerebro para pagarés
                        json_respuesta_pagare = generar_pagare_ia(instruccion, texto_de_ejemplo, modelo_seleccionado)
                    
                    if json_respuesta_pagare:
                        with st.spinner("🛠️ Ensamblando el archivo .docx del pagaré..."):
                            # Usamos la nueva función del ensamblador para pagarés
                            buffer_docx_pagare = ensamblar_pagare_en_memoria(json_respuesta_pagare, TEMPLATE_FILE_PAGARE)
                        
                        if buffer_docx_pagare:
                            st.success("¡Tu Pagaré con Tabla está listo para descargar!")
                            st.download_button(
                                label="📥 Descargar Pagaré Generado",
                                data=buffer_docx_pagare,
                                file_name="PAGARE_GENERADO.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

# --- PÁGINA DE ADMIN (V2.3 - CRUD) ---
def mostrar_pagina_admin():
    
    st.title("Panel de Administración de Usuarios")

    # 1. Crear Usuario (CREATE)
    st.subheader("1. Crear Nuevo Usuario")
    with st.form("create_user_form", clear_on_submit=True):
        new_email = st.text_input("Email del nuevo usuario")
        new_password = st.text_input("Contraseña temporal para el usuario", type="password")
        create_button = st.form_submit_button("Crear Usuario")
        
        # --- ¡ESTE ES EL CÓDIGO NUEVO Y CORREGIDO (V2.4)! ---
    if create_button:
        if new_email and new_password:
            try:
                # ¡EL ARREGLO ESTÁ AQUÍ! 
                # Pasamos los datos dentro de un diccionario "attributes"
                user = supabase_admin.auth.admin.create_user(
                    attributes={
                        'email': new_email,
                        'password': new_password,
                        'email_confirm': True 
                    }
                )
                st.success(f"¡Éxito! Usuario creado: {user.user.email}")
            except Exception as e:
                st.error(f"Error al crear usuario: {e}")
        else:
            st.warning("Por favor, llena ambos campos.")
    # --- FIN DEL CÓDIGO NUEVO ---

    st.markdown("---")
    
    # 2. Listar Usuarios (READ)
    st.subheader("2. Usuarios Actuales")
    # --- ¡ESTE ES EL CÓDIGO NUEVO Y CORREGIDO (V2.5)! ---
    if st.button("Cargar Lista de Usuarios"):
        try:
            # Esta línea sigue igual
            users = supabase_admin.auth.admin.list_users()

            # ¡EL ARREGLO ESTÁ AQUÍ! 
            # Iteramos sobre 'users' directamente, no 'users.users'
            st.dataframe(
                [{"email": u.email, "created_at": u.created_at, "id": u.id} for u in users]
            )
        except Exception as e:
            st.error(f"Error al listar usuarios: {e}")
    # --- FIN DEL CÓDIGO NUEVO ---

    # 3. Borrar Usuario (DELETE)
    st.markdown("---")
    st.subheader("3. Borrar Usuario (¡Peligro!)")
    with st.form("delete_user_form", clear_on_submit=True):
        user_id_to_delete = st.text_input("Pega el 'ID' del usuario a borrar (de la tabla de arriba)")
        delete_button = st.form_submit_button("Borrar Usuario Permanentemente")
        
        if delete_button and user_id_to_delete:
            try:
                # ¡Usamos el cliente ADMIN!
                # ¡EL ARREGLO! El parámetro es 'id', no 'user_id'
                supabase_admin.auth.admin.delete_user(id=user_id_to_delete) 
                st.success(f"¡Éxito! Usuario {user_id_to_delete} borrado.")
            except Exception as e:
                st.error(f"Error al borrar usuario: {e}")

# --- ¡NUEVA FUNCIÓN PARA EL CHATBOT! ---

# --- CEREBRO DEL CHATBOT (SIMPLIFICADO) ---
def llamar_chat_ia(historial_mensajes, modelo_ia):
    """
    Función 'Cerebro' simple para el chatbot.
    Usa el endpoint v1/chat/completions (perfecto para gpt-4o-mini).
    """
    try:
        completion = client.chat.completions.create(
            model=modelo_ia, 
            messages=historial_mensajes
        )
        return completion.choices[0].message.content
    except Exception as e:
        st.error(f"Error al llamar a la API del chat: {e}")
        return None

# --- PÁGINA DEL CHATBOT (V 1.7 - ANALIZADOR CON RAZONAMIENTO) ---
def mostrar_pagina_chatbot():
    """
    Contiene la lógica para el chatbot (V1.7),
    CON PERMISO DE RAZONAR SOBRE EL DOCUMENTO.
    """
    
    st.warning("⚠️ **Modo Demo:** Este chatbot es una IA generativa (gpt-4o). Sus respuestas pueden ser imprecisas. Usa tu criterio profesional.")
    
    # Forzar el modelo rápido
    modelo_chatbot = "gpt-4o" # ¡Cambiamos a gpt-4o (no el mini)!

    # --- LÓGICA DE CARGA DE DOCUMENTO ---
    uploaded_doc = st.file_uploader(
        "Sube un .docx o .pdf para analizarlo:",
        type=["docx", "pdf"],
        key="chatbot_file_uploader"
    )
    
    if uploaded_doc is not None:
        if "documento_analizado" not in st.session_state or st.session_state.documento_analizado != uploaded_doc.name:
            with st.spinner(f"Analizando '{uploaded_doc.name}'..."):
                contexto_documento = extraer_texto_del_documento(uploaded_doc)
                if contexto_documento:
                    
                    # --- ¡EL NUEVO PROMPT DE SISTEMA (V1.7)! ---
                    st.session_state.chat_history = [
                        {
                            "role": "system",
                            "content": f"""
                            Eres un Asistente Legal experto en análisis de documentos. Te han pasado el siguiente documento como contexto.

                            TUS REGLAS:
                            1.  **Basa tus respuestas SIEMPRE en el documento.**
                            2.  **SÍ PUEDES USAR TU PROPIO RAZONAMIENTO** y habilidades matemáticas para analizar el contenido (ej. verificar cálculos, resumir, identificar riesgos).
                            3.  Si te preguntan algo que NO está en el documento y NO se relaciona con él (ej. el clima, política, "quién es el presidente"), DEBES responder: "Mi función es analizar el documento proporcionado."
                            
                            ---
                            CONTEXTO DEL DOCUMENTO (Nombre: {uploaded_doc.name}):
                            {contexto_documento}
                            ---
                            """
                        },
                        {
                            "role": "assistant",
                            "content": f"He leído y analizado el documento '{uploaded_doc.name}'. ¿Qué necesitas que revise o calcule?"
                        }
                    ]
                    # --- FIN DEL NUEVO PROMPT ---
                    
                    st.session_state.documento_analizado = uploaded_doc.name 
                    st.success("Documento analizado. ¡Puedes empezar a chatear con él!")
                else:
                    st.error("No se pudo extraer texto del documento.")

    # --- INICIALIZACIÓN DE CHAT (SI NO HAY DOC) ---
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = [
            {"role": "assistant", "content": "Hola, soy el asistente legal. Sube un documento para analizar o hazme una pregunta general."}
        ]

    # --- 1. PINTAR EL HISTORIAL ANTIGUO (CON FILTRO) ---
    for message in st.session_state.chat_history:
        if message["role"] != "system": # Oculta el prompt del sistema
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

    # --- 2. PEDIR NUEVO INPUT (Lógica de V1.7) ---
    if prompt := st.chat_input("Escribe tu pregunta aquí..."):
        
        # Guardar mensaje del usuario en la memoria
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        
        # (¡No pintamos aquí! Usamos st.rerun())

        # Generar respuesta de la IA
        with st.spinner("Analizando..."):
            respuesta_ia = llamar_chat_ia(
                st.session_state.chat_history, 
                modelo_chatbot
            )
            
            if respuesta_ia:
                # Guardar respuesta de la IA en la memoria
                st.session_state.chat_history.append(
                    {"role": "assistant", "content": respuesta_ia}
                )
            else:
                st.error("No se pudo obtener respuesta de la IA.")
        
        # Forzar un refresco de la página
        st.rerun()

# --- LOGIN (V2.2 con Supabase) ---
def mostrar_login():
    col1, col2, col3 = st.columns([1, 2, 1]) 
    with col2:
        st.image("logoOscuro.png")

    st.title("Asistente Legal 🔒")
    
    st.markdown("Por favor, inicia sesión para continuar.")

    # --- 1. Formulario de Login ---
    with st.form("login_form"):
        email = st.text_input("Email")
        password = st.text_input("Contraseña", type="password")
        submitted = st.form_submit_button("Iniciar Sesión")

        if submitted:
            try:
                # Usamos el cliente ANÓNIMO para el login
                session = supabase_anon.auth.sign_in_with_password({
                    "email": email,
                    "password": password
                })
                st.session_state['authenticated'] = True
                st.session_state['user_email'] = session.user.email
                st.rerun() 
            except Exception as e:
                st.error(f"Error de autenticación: Verifique su email o contraseña.")

def logout():
    st.session_state['authenticated'] = False
    st.session_state.pop('user_email', None) 
    st.session_state.pop('chat_history', None)
    st.session_state.pop('documento_analizado', None)

# --- 5. LÓGICA PRINCIPAL (EL "SWITCH") ---

# Inicializar el estado de sesión
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

# Mostrar la página correspondiente
if st.session_state['authenticated']:
    # Renombramos la función original
    mostrar_app_principal() 
else:
    mostrar_login()